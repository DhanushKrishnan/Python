import docx

def is_question(paragraph):
    # Customize this function to identify questions based on specific criteria.
    return paragraph.text.endswith('?')

def change_questions_format(input_file, output_file, font_family, font_size):
    doc = docx.Document(input_file)

    # Convert the first line to title size
    if doc.paragraphs:
        first_paragraph = doc.paragraphs[0]
        for run in first_paragraph.runs:
            run.font.size = docx.shared.Pt(font_size * 2)
            run.bold = True
            run.font.color.rgb = docx.shared.RGBColor(*font_color)

    
    for paragraph in doc.paragraphs:
        if is_question(paragraph):
            for run in paragraph.runs:
                run.bold = True
                run.font.name = font_family
                run.font.size = docx.shared.Pt(font_size)

    doc.save(output_file)

# Replace 'path' with the path of your input document

input_file = 'path'
output_file = 'path'
font_family = 'Libre Baskerville'  # Customize the font family as desired
font_size = 12        # Customize the font size as desired
font_color = (65, 105, 225)   #customize font color of Title

change_questions_format(input_file, output_file, font_family, font_size)
